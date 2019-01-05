import docx

a = '关键词'
#文件路径
doc5 = docx.Document('E:\\spark_dataset\\2005.docx')
#2005
f5 = []
for i in range(len(doc5.paragraphs)):
    if a in doc5.paragraphs[i].text:
        key_C = doc5.paragraphs[i].text[4:]
        f5.append(key_C)
t1 = open('E:\\a\\2005.txt', 'w')
for item in f5:
    t1.write(item+'\n')
t1.close()


doc6 = docx.Document('E:\\spark_dataset\\2006.docx')
doc7 = docx.Document('E:\\spark_dataset\\2007.docx')
doc8 = docx.Document('E:\\spark_dataset\\2008.docx')
doc9 = docx.Document('E:\\spark_dataset\\2009.docx')
doc10 = docx.Document('E:\\spark_dataset\\2010.docx')
doc11 = docx.Document('E:\\spark_dataset\\2011.docx')
doc12 = docx.Document('E:\\spark_dataset\\2012.docx')
doc13 = docx.Document('E:\\spark_dataset\\2013.docx')
doc14 = docx.Document('E:\\spark_dataset\\2014.docx')
doc15 = docx.Document('E:\\spark_dataset\\2015.docx')
doc16 = docx.Document('E:\\spark_dataset\\2016.docx')
doc17 = docx.Document('E:\\spark_dataset\\2017.docx')
doc18 = docx.Document('E:\\spark_dataset\\2018.docx')

print('提取各年的关键词到E盘a文件夹下')




#2006
f6 = []
for i in range(len(doc6.paragraphs)):
    if a in doc6.paragraphs[i].text:
        key_C = doc6.paragraphs[i].text[4:]
        f6.append(key_C)
t2 = open('E:\\a\\2006.txt', 'w')
for item in f6:
    t2.write(item+'\n')
t2.close()

#2007
f7 = []
for i in range(len(doc7.paragraphs)):
    if a in doc7.paragraphs[i].text:
        key_C = doc7.paragraphs[i].text[4:]
        f7.append(key_C)
t3 = open('E:\\a\\2007.txt', 'w')
for item in f7:
    t3.write(item+'\n')
t3.close()

#2008
f8 = []
for i in range(len(doc8.paragraphs)):
    if a in doc8.paragraphs[i].text:
        key_C = doc8.paragraphs[i].text[4:]
        f8.append(key_C)
t4 = open('E:\\a\\2008.txt', 'w')
for item in f8:
    t4.write(item+'\n')
t4.close()

2009
f9 = []
for i in range(len(doc9.paragraphs)):
    if a in doc9.paragraphs[i].text:
        key_C = doc9.paragraphs[i].text[4:]
        # print(key_C)
        f9.append(key_C)
t5 = open('E:\\a\\2009.txt', 'w',encoding='utf-8')
for item in f9:
    t5.write(item+'\n')
t5.close()

#2010
f10 = []
for i in range(len(doc10.paragraphs)):
    if a in doc10.paragraphs[i].text:
        key_C = doc10.paragraphs[i].text[4:]
        f10.append(key_C)
t6 = open('E:\\a\\2010.txt', 'w')
for item in f10:
    t6.write(item+'\n')
t6.close()

#2011
f11 = []
for i in range(len(doc11.paragraphs)):
    if a in doc11.paragraphs[i].text:
        key_C = doc11.paragraphs[i].text[4:]
        f11.append(key_C)
t7 = open('E:\\a\\2011.txt', 'w')
for item in f11:
    t7.write(item+'\n')
t7.close()

#2012
f12 = []
for i in range(len(doc12.paragraphs)):
    if a in doc12.paragraphs[i].text:
        key_C = doc12.paragraphs[i].text[4:]
        f12.append(key_C)
t8 = open('E:\\a\\2012.txt', 'w')
for item in f12:
    t8.write(item+'\n')
t8.close()

#2013
f13 = []
for i in range(len(doc13.paragraphs)):
    if a in doc13.paragraphs[i].text:
        key_C = doc13.paragraphs[i].text[4:]
        f13.append(key_C)
t9 = open('E:\\a\\2013.txt', 'w')
for item in f13:
    t9.write(item+'\n')
t9.close()

#2014
f14 = []
for i in range(len(doc14.paragraphs)):
    if a in doc14.paragraphs[i].text:
        key_C = doc14.paragraphs[i].text[4:]
        f14.append(key_C)
t10 = open('E:\\a\\2014.txt', 'w')
for item in f14:
    t10.write(item+'\n')
t10.close()

#2015
f15 = []
for i in range(len(doc15.paragraphs)):
    if a in doc15.paragraphs[i].text:
        key_C = doc15.paragraphs[i].text[4:]
        f15.append(key_C)
t11 = open('E:\\a\\2015.txt', 'w')
for item in f15:
    t11.write(item+'\n')
t11.close()

#2016
f16 = []
for i in range(len(doc16.paragraphs)):
    if a in doc16.paragraphs[i].text:
        key_C = doc16.paragraphs[i].text[4:]
        f16.append(key_C)
t12 = open('E:\\a\\2016.txt', 'w')
for item in f16:
    t12.write(item+'\n')
t12.close()

#2017
f17 = []
for i in range(len(doc17.paragraphs)):
    if a in doc17.paragraphs[i].text:
        key_C = doc17.paragraphs[i].text[4:]
        f17.append(key_C)
t13 = open('E:\\a\\2017.txt', 'w',encoding='utf-8')
for item in f17:
    t13.write(item+'\n')
t13.close()

#2018
f18 = []
for i in range(len(doc18.paragraphs)):
    if a in doc18.paragraphs[i].text:
        key_C = doc18.paragraphs[i].text[4:]
        f18.append(key_C)
t14 = open('E:\\a\\2018.txt', 'w')
for item in f18:
    t14.write(item+'\n')
t14.close()